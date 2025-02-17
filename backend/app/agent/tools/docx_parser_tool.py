from langchain.tools import BaseTool
from docx import Document
import io
import re
from .utils import chunk_text_by_paragraphs

class DOCXParserTool(BaseTool):
    name: str = "docx_parser"
    description: str = "Extracts text content from a docx file. Input should be bytes object. Output is the extracted text content."

    def _run(self, content: bytes, max_chunk_size: int = 1500, min_chunk_size: int = 500) -> str:
        try:
            doc = Document(io.BytesIO(content))

            # Extract text from paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join([cell.text for cell in row.cells])

            # Remove invalid unicode characters
            cleaned_text = re.sub(r"/uni[0-9a-fA-F]+", "", text).strip()

            # Chunk the extracted text
            chunks = chunk_text_by_paragraphs(cleaned_text, max_chunk_size, min_chunk_size)
            return chunks
        
        except Exception as e:
            return f"Error processing Docx: {str(e)}"